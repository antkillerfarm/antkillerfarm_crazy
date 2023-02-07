# -*- coding: utf-8 -*-

# ref: https://www.cnblogs.com/techlynn/p/9736351.html

import random
from datetime import datetime
from docx import Document
from docx.shared import Pt

def add_test(sum_min, sum_max, add_max, count):
    questions = []
    count_temp = 0

    while True:
        i = random.randrange(0, add_max + 1)  # 随机生成 第一个加数
        j = random.randrange(sum_min, sum_max + 1)  # 随机生成 和
        l = j - i  # 第二个加数
        if l >= 0 and l <= add_max:
            questions.append((i, l, j))
            count_temp += 1
            if count_temp >= count:
                break

    return questions


def resort(quiz):
    rng_index = 2
    flag_addsub = random.randint(0, 1)
    if flag_addsub:
        str_temp = (str(quiz[0]) if rng_index != 0 else '(   )') + ' + ' \
            + (str(quiz[1]) if rng_index != 1 else '(   )') \
            + ' = ' \
            + (str(quiz[2]) if rng_index != 2 else '(   )')
    else:
        str_temp = (str(quiz[2]) if rng_index != 0 else '(   )') + ' - ' \
            + (str(quiz[0]) if rng_index != 1 else '(   )') \
            + ' = ' \
            + (str(quiz[1]) if rng_index != 2 else '(   )')
    return str_temp


def main():
    sum_min, sum_max, add_max, count = 11, 18, 9, 200
    quizs = add_test(sum_min, sum_max, add_max, count)
    quizs0 = []
    text = ''
    for id, quiz in enumerate(quizs):
        if (id % 10 == 0 and id != 0):
            quizs0.append(text)
            text = ''
        text += resort(quiz)
        if (id % 10 != 9):
            text += "\n"
    quizs0.append(text)

    document = Document()
    cols = 4
    table = document.add_table(rows=0, cols=cols)
    for index, quiz0 in enumerate(quizs0):
        if (index % cols == 0):
            row_cells = table.add_row().cells
        row_cells[index % cols].text = quiz0
        row_cells[index % cols].paragraphs[0].runs[0].font.size = Pt(15)
    document.save('demo_18.docx')


if __name__ == '__main__':
    main()
